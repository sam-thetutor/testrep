#!/usr/bin/env python3
"""
Magnus Client Intake Form Generator
- Save drafts in Word format (.docx) for editing
- Generate final reports in PDF format

INSTALLATION REQUIRED:
pip install reportlab python-docx
"""

import os
import sys
import traceback

# Check for required packages
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import inch
except ImportError:
    print("ERROR: ReportLab is not installed. Please run: pip install reportlab")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx is not installed. Please run: pip install python-docx")
    sys.exit(1)

def save_draft_word(form_data, output_path):
    """Save form data as a Word document draft"""
    try:
        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading('Magnus Client Intake Form', 0)
        doc.add_paragraph()
        
        # Personal Information
        doc.add_heading('Personal Information', level=1)
        doc.add_paragraph(f"Full Name: {form_data.get('full_name', '[Not provided]')}")
        doc.add_paragraph(f"Date of Birth: {form_data.get('dob', '[Not provided]')}")
        doc.add_paragraph(f"Social Security Number: {form_data.get('ssn', '[Not provided]')}")
        doc.add_paragraph(f"Citizenship: {form_data.get('citizenship', '[Not provided]')}")
        doc.add_paragraph(f"Marital Status: {form_data.get('marital_status', '[Not provided]')}")
        doc.add_paragraph()
        
        # Contact Information
        doc.add_heading('Contact Information', level=1)
        doc.add_paragraph(f"Residential Address: {form_data.get('residential_address', '[Not provided]')}")
        
        # Mailing Address (if different)
        if form_data.get("mailing_address_different", False):
            doc.add_paragraph(f"Mailing Address: {form_data.get('mailing_address', '[Not provided]')}")
        
        doc.add_paragraph(f"Home Phone: {form_data.get('home_phone', '[Not provided]')}")
        doc.add_paragraph(f"Work Phone: {form_data.get('work_phone', '[Not provided]')}")
        doc.add_paragraph(f"Mobile Phone: {form_data.get('mobile_phone', '[Not provided]')}")
        doc.add_paragraph(f"Email: {form_data.get('email', '[Not provided]')}")
        doc.add_paragraph()
        
        # Employment Information
        doc.add_heading('Employment Information', level=1)
        doc.add_paragraph(f"Employment Status: {form_data.get('employment_status', '[Not provided]')}")
        doc.add_paragraph(f"Employer Name: {form_data.get('employer_name', '[Not provided]')}")
        doc.add_paragraph(f"Occupation/Title: {form_data.get('occupation', '[Not provided]')}")
        doc.add_paragraph(f"Years Employed: {form_data.get('years_employed', '[Not provided]')}")
        doc.add_paragraph(f"Annual Income: {form_data.get('annual_income', '[Not provided]')}")
        doc.add_paragraph(f"Employer Address: {form_data.get('employer_address', '[Not provided]')}")
        
        # Retirement Information (if applicable)
        if form_data.get("employment_status") == "Retired":
            doc.add_paragraph()
            doc.add_heading('Retirement Information', level=1)
            doc.add_paragraph(f"Former Employer: {form_data.get('former_employer', '[Not provided]')}")
            doc.add_paragraph(f"Source of Income: {form_data.get('income_source', '[Not provided]')}")
        
        doc.add_paragraph()
        
        # Spouse Information (if applicable)
        if form_data.get("spouse_applicable", False):
            doc.add_heading('Spouse/Partner Information', level=1)
            doc.add_paragraph(f"Spouse Full Name: {form_data.get('spouse_full_name', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Date of Birth: {form_data.get('spouse_dob', '[Not provided]')}")
            doc.add_paragraph(f"Spouse SSN: {form_data.get('spouse_ssn', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Employment Status: {form_data.get('spouse_employment_status', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Employer Name: {form_data.get('spouse_employer_name', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Occupation/Title: {form_data.get('spouse_occupation', '[Not provided]')}")
            doc.add_paragraph()
        
        # Dependents Information
        if isinstance(form_data.get("dependents"), list):
            doc.add_heading('Dependents Information', level=1)
            if form_data["dependents"]:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Name'
                header_cells[1].text = 'Date of Birth'
                header_cells[2].text = 'Relationship'
                
                for dep in form_data["dependents"]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = dep.get("name", "[Not provided]")
                    row_cells[1].text = dep.get("dob", "[Not provided]")
                    row_cells[2].text = dep.get("relationship", "[Not provided]")
            else:
                doc.add_paragraph("No dependents specified")
            doc.add_paragraph()
        
        # Beneficiaries Information
        if isinstance(form_data.get("beneficiaries"), list):
            doc.add_heading('Beneficiaries Information', level=1)
            if form_data["beneficiaries"]:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Name'
                header_cells[1].text = 'Date of Birth'
                header_cells[2].text = 'Relationship'
                header_cells[3].text = 'Percentage'
                
                for ben in form_data["beneficiaries"]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = ben.get("name", "[Not provided]")
                    row_cells[1].text = ben.get("dob", "[Not provided]")
                    row_cells[2].text = ben.get("relationship", "[Not provided]")
                    percentage = ben.get('percentage', '')
                    row_cells[3].text = f"{percentage}%" if percentage else "[Not provided]"
            else:
                doc.add_paragraph("No beneficiaries specified")
            doc.add_paragraph()
        
        # Assets & Investment Experience
        doc.add_heading('Assets & Investment Experience', level=1)
        doc.add_paragraph(f"Net Worth (excluding primary home): {form_data.get('net_worth', '[Not provided]')}")
        doc.add_paragraph(f"Liquid Net Worth: {form_data.get('liquid_net_worth', '[Not provided]')}")
        doc.add_paragraph(f"Assets Held Away: {form_data.get('assets_held_away', '[Not provided]')}")
        doc.add_paragraph(f"Annual Income: {form_data.get('annual_income', '[Not provided]')}")
        doc.add_paragraph(f"Tax Bracket: {form_data.get('tax_bracket', '[Not provided]')}")
        doc.add_paragraph(f"Education Status: {form_data.get('education_status', '[Not provided]')}")
        doc.add_paragraph(f"Risk Tolerance: {form_data.get('risk_tolerance', '[Not provided]')}")
        doc.add_paragraph(f"Investment Objectives: {form_data.get('investment_objectives', '[Not provided]')}")
        
        # Asset Breakdown (if included)
        if form_data.get("include_breakdown", False):
            doc.add_paragraph()
            doc.add_heading('Asset Breakdown', level=1)
            asset_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures", "Short-Term", "Other"]
            for asset_type in asset_types:
                field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_')}"
                value = form_data.get(field_name, "[Not provided]")
                doc.add_paragraph(f"{asset_type}: {value}%")
        
        # Investment Experience
        doc.add_paragraph()
        doc.add_heading('Investment Experience', level=1)
        experience_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures"]
        for exp_type in experience_types:
            year_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_year"
            level_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_level"
            
            year = form_data.get(year_field, "[Not provided]")
            level = form_data.get(level_field, "[Not provided]")
            
            doc.add_paragraph(f"{exp_type}:")
            doc.add_paragraph(f"  Year Started: {year}")
            doc.add_paragraph(f"  Experience Level: {level}")
            doc.add_paragraph()
        
        # Outside Broker Information
        if form_data.get("has_outside_broker", False):
            doc.add_heading('Outside Broker Information', level=1)
            doc.add_paragraph(f"Broker Firm Name: {form_data.get('outside_broker_name', '[Not provided]')}")
            doc.add_paragraph(f"Account Number: {form_data.get('outside_broker_account', '[Not provided]')}")
            doc.add_paragraph(f"Account Type: {form_data.get('outside_broker_account_type', '[Not provided]')}")
            doc.add_paragraph()
        
        # Trusted Contact Information
        doc.add_heading('Trusted Contact Information', level=1)
        doc.add_paragraph(f"Full Name: {form_data.get('trusted_full_name', '[Not provided]')}")
        doc.add_paragraph(f"Relationship: {form_data.get('trusted_relationship', '[Not provided]')}")
        doc.add_paragraph(f"Phone Number: {form_data.get('trusted_phone', '[Not provided]')}")
        doc.add_paragraph(f"Email Address: {form_data.get('trusted_email', '[Not provided]')}")
        doc.add_paragraph()
        
        # Regulatory Consent
        doc.add_heading('Regulatory Consent', level=1)
        electronic_consent = "Yes" if form_data.get("electronic_regulatory_yes", False) else "No"
        doc.add_paragraph(f"Electronic Delivery Consent: {electronic_consent}")
        
        # Save document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error saving Word document: {str(e)}")
        traceback.print_exc()
        return False

def generate_pdf_report(form_data, output_path):
    """Generate PDF report from form data"""
    try:
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = styles['Heading1']
        subtitle_style = styles['Heading2']
        normal_style = styles['Normal']
        
        # Build content
        content = []
        
        # Title
        content.append(Paragraph("Magnus Client Intake Form", title_style))
        content.append(Spacer(1, 12))
        
        # Personal Information
        content.append(Paragraph("Personal Information", subtitle_style))
        content.append(Spacer(1, 6))
        
        personal_info = [
            ("Full Name", form_data.get("full_name", "[Not provided]")),
            ("Date of Birth", form_data.get("dob", "[Not provided]")),
            ("Social Security Number", form_data.get("ssn", "[Not provided]")),
            ("Citizenship", form_data.get("citizenship", "[Not provided]")),
            ("Marital Status", form_data.get("marital_status", "[Not provided]"))
        ]
        
        for label, value in personal_info:
            content.append(Paragraph(f"{label}: {value}", normal_style))
        
        content.append(Spacer(1, 12))
        
        # Contact Information
        content.append(Paragraph("Contact Information", subtitle_style))
        content.append(Spacer(1, 6))
        
        # Residential Address
        content.append(Paragraph(f"Residential Address: {form_data.get('residential_address', '[Not provided]')}", normal_style))
        
        # Mailing Address (if different)
        if form_data.get("mailing_address_different", False):
            content.append(Paragraph(f"Mailing Address: {form_data.get('mailing_address', '[Not provided]')}", normal_style))
        
        contact_info = [
            ("Home Phone", form_data.get("home_phone", "[Not provided]")),
            ("Work Phone", form_data.get("work_phone", "[Not provided]")),
            ("Mobile Phone", form_data.get("mobile_phone", "[Not provided]")),
            ("Email", form_data.get("email", "[Not provided]"))
        ]
        
        for label, value in contact_info:
            content.append(Paragraph(f"{label}: {value}", normal_style))
        
        content.append(Spacer(1, 12))
        
        # Employment Information
        content.append(Paragraph("Employment Information", subtitle_style))
        content.append(Spacer(1, 6))
        
        employment_info = [
            ("Employment Status", form_data.get("employment_status", "[Not provided]")),
            ("Employer Name", form_data.get("employer_name", "[Not provided]")),
            ("Occupation/Title", form_data.get("occupation", "[Not provided]")),
            ("Years Employed", str(form_data.get("years_employed", "[Not provided]"))),
            ("Annual Income", form_data.get("annual_income", "[Not provided]"))
        ]
        
        for label, value in employment_info:
            content.append(Paragraph(f"{label}: {value}", normal_style))
        
        # Employer Address
        content.append(Paragraph(f"Employer Address: {form_data.get('employer_address', '[Not provided]')}", normal_style))
        
        # Retirement Information (if applicable)
        if form_data.get("employment_status") == "Retired":
            content.append(Spacer(1, 12))
            content.append(Paragraph("Retirement Information", subtitle_style))
            content.append(Spacer(1, 6))
            
            retirement_info = [
                ("Former Employer", form_data.get("former_employer", "[Not provided]")),
                ("Source of Income", form_data.get("income_source", "[Not provided]"))
            ]
            
            for label, value in retirement_info:
                content.append(Paragraph(f"{label}: {value}", normal_style))
        
        content.append(Spacer(1, 12))
        
        # Spouse Information (if applicable)
        if form_data.get("spouse_applicable", False):
            content.append(Paragraph("Spouse/Partner Information", subtitle_style))
            content.append(Spacer(1, 6))
            
            spouse_info = [
                ("Spouse Full Name", form_data.get("spouse_full_name", "[Not provided]")),
                ("Spouse Date of Birth", form_data.get("spouse_dob", "[Not provided]")),
                ("Spouse SSN", form_data.get("spouse_ssn", "[Not provided]")),
                ("Spouse Employment Status", form_data.get("spouse_employment_status", "[Not provided]")),
                ("Spouse Employer Name", form_data.get("spouse_employer_name", "[Not provided]")),
                ("Spouse Occupation/Title", form_data.get("spouse_occupation", "[Not provided]"))
            ]
            
            for label, value in spouse_info:
                content.append(Paragraph(f"{label}: {value}", normal_style))
            
            content.append(Spacer(1, 12))
        
        # Dependents Information
        if isinstance(form_data.get("dependents"), list):
            content.append(Paragraph("Dependents Information", subtitle_style))
            content.append(Spacer(1, 6))
            
            if form_data["dependents"]:
                # Create table for dependents
                data = [['Name', 'Date of Birth', 'Relationship']]
                for dep in form_data["dependents"]:
                    data.append([
                        dep.get("name", "[Not provided]"),
                        dep.get("dob", "[Not provided]"),
                        dep.get("relationship", "[Not provided]")
                    ])
                
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                content.append(table)
            else:
                content.append(Paragraph("No dependents specified", normal_style))
            
            content.append(Spacer(1, 12))
        
        # Beneficiaries Information
        if isinstance(form_data.get("beneficiaries"), list):
            content.append(Paragraph("Beneficiaries Information", subtitle_style))
            content.append(Spacer(1, 6))
            
            if form_data["beneficiaries"]:
                # Create table for beneficiaries
                data = [['Name', 'Date of Birth', 'Relationship', 'Percentage']]
                for ben in form_data["beneficiaries"]:
                    percentage = ben.get('percentage', '')
                    data.append([
                        ben.get("name", "[Not provided]"),
                        ben.get("dob", "[Not provided]"),
                        ben.get("relationship", "[Not provided]"),
                        f"{percentage}%" if percentage else "[Not provided]"
                    ])
                
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                content.append(table)
            else:
                content.append(Paragraph("No beneficiaries specified", normal_style))
            
            content.append(Spacer(1, 12))
        
        # Assets & Investment Experience
        content.append(Paragraph("Assets & Investment Experience", subtitle_style))
        content.append(Spacer(1, 6))
        
        assets_info = [
            ("Net Worth (excluding primary home)", form_data.get("net_worth", "[Not provided]")),
            ("Liquid Net Worth", form_data.get("liquid_net_worth", "[Not provided]")),
            ("Assets Held Away", form_data.get("assets_held_away", "[Not provided]")),
            ("Annual Income", form_data.get("annual_income", "[Not provided]")),
            ("Tax Bracket", form_data.get("tax_bracket", "[Not provided]")),
            ("Education Status", form_data.get("education_status", "[Not provided]")),
            ("Risk Tolerance", form_data.get("risk_tolerance", "[Not provided]")),
            ("Investment Objectives", form_data.get("investment_objectives", "[Not provided]"))
        ]
        
        for label, value in assets_info:
            content.append(Paragraph(f"{label}: {value}", normal_style))
        
        # Asset Breakdown (if included)
        if form_data.get("include_breakdown", False):
            content.append(Spacer(1, 12))
            content.append(Paragraph("Asset Breakdown", subtitle_style))
            content.append(Spacer(1, 6))
            
            asset_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures", "Short-Term", "Other"]
            for asset_type in asset_types:
                field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_')}"
                value = form_data.get(field_name, "[Not provided]")
                content.append(Paragraph(f"{asset_type}: {value}%", normal_style))
        
        # Investment Experience
        content.append(Spacer(1, 12))
        content.append(Paragraph("Investment Experience", subtitle_style))
        content.append(Spacer(1, 6))
        
        experience_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures"]
        for exp_type in experience_types:
            year_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_year"
            level_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_level"
            
            year = form_data.get(year_field, "[Not provided]")
            level = form_data.get(level_field, "[Not provided]")
            
            content.append(Paragraph(f"{exp_type}:", normal_style))
            content.append(Paragraph(f"  Year Started: {year}", normal_style))
            content.append(Paragraph(f"  Experience Level: {level}", normal_style))
            content.append(Spacer(1, 6))
        
        # Outside Broker Information
        if form_data.get("has_outside_broker", False):
            content.append(Spacer(1, 12))
            content.append(Paragraph("Outside Broker Information", subtitle_style))
            content.append(Spacer(1, 6))
            
            broker_info = [
                ("Broker Firm Name", form_data.get("outside_broker_name", "[Not provided]")),
                ("Account Number", form_data.get("outside_broker_account", "[Not provided]")),
                ("Account Type", form_data.get("outside_broker_account_type", "[Not provided]"))
            ]
            
            for label, value in broker_info:
                content.append(Paragraph(f"{label}: {value}", normal_style))
        
        # Trusted Contact Information
        content.append(Spacer(1, 12))
        content.append(Paragraph("Trusted Contact Information", subtitle_style))
        content.append(Spacer(1, 6))
        
        trusted_info = [
            ("Full Name", form_data.get("trusted_full_name", "[Not provided]")),
            ("Relationship", form_data.get("trusted_relationship", "[Not provided]")),
            ("Phone Number", form_data.get("trusted_phone", "[Not provided]")),
            ("Email Address", form_data.get("trusted_email", "[Not provided]"))
        ]
        
        for label, value in trusted_info:
            content.append(Paragraph(f"{label}: {value}", normal_style))
        
        # Regulatory Consent
        content.append(Spacer(1, 12))
        content.append(Paragraph("Regulatory Consent", subtitle_style))
        content.append(Spacer(1, 6))
        
        electronic_consent = "Yes" if form_data.get("electronic_regulatory_yes", False) else "No"
        content.append(Paragraph(f"Electronic Delivery Consent: {electronic_consent}", normal_style))
        
        # Build the PDF
        doc.build(content)
        return True
        
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        traceback.print_exc()
        return False

# Alias for backward compatibility with main_enhanced.py
generate_pdf_from_data = generate_pdf_report