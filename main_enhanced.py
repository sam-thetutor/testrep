#!/usr/bin/env python3
"""
Magnus Client Intake Form - Enhanced Version 2.2
Professional client intake form for financial services with comprehensive validation,
security features, and PDF generation capabilities.
"""

import sys
import os
import json
import tempfile
from datetime import datetime
from typing import Dict, Any, List

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QLabel, QLineEdit, QPushButton, QStackedWidget, QFrame, 
    QComboBox, QDateEdit, QTextEdit, QCheckBox, QRadioButton,
    QButtonGroup, QSpinBox, QGroupBox, QScrollArea, QMessageBox,
    QProgressBar, QFileDialog, 
)
import traceback
from PyQt6.QtCore import Qt, QDate, QTimer, pyqtSignal
from PyQt6.QtGui import QFont, QPixmap, QIcon
import docx
import subprocess
import platform
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Import custom modules
from validation import form_validator
from security import DataSecurity
from pdf_generator_reportlab import generate_pdf_from_data


class EnhancedLineEdit(QLineEdit):
    """Enhanced QLineEdit with validation feedback"""
    
    def __init__(self, field_name: str, parent=None):
        super().__init__(parent)
        self.field_name = field_name
        self.setStyleSheet("""
            QLineEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
            QLineEdit:focus {
                border-color: #4CAF50;
            }
        """)
        
    def validate_field(self) -> bool:
        """Validate field content and update styling"""
        text = self.text().strip()
        
        # Basic validation based on field name
        if "email" in self.field_name.lower():
            valid = form_validator.validate_email(self.field_name, text)
        elif "ssn" in self.field_name.lower():
            valid = form_validator.validate_ssn(self.field_name, text)
        elif "phone" in self.field_name.lower():
            valid = form_validator.validate_phone(self.field_name, text)
        else:
            valid = len(text) > 0 if text else True
        
        # Update styling based on validation
        if text and not valid:
            self.setStyleSheet("""
                QLineEdit {
                    padding: 8px;
                    border: 2px solid #f44336;
                    border-radius: 4px;
                    font-size: 12px;
                    background-color: #ffebee;
                }
            """)
        elif text and valid:
            self.setStyleSheet("""
                QLineEdit {
                    padding: 8px;
                    border: 2px solid #4CAF50;
                    border-radius: 4px;
                    font-size: 12px;
                    background-color: #e8f5e8;
                }
            """)
        else:
            self.setStyleSheet("""
                QLineEdit {
                    padding: 8px;
                    border: 2px solid #ddd;
                    border-radius: 4px;
                    font-size: 12px;
                }
            """)
        
        return valid


class MagnusClientIntakeForm(QMainWindow):
    """Main application window for Magnus Client Intake Form"""
    
    def __init__(self):
        super().__init__()
        self.security_manager = DataSecurity()
        self.form_data = {}
        self.current_page = 0
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save_data)
        self.auto_save_timer.start(30000)  # Auto-save every 30 seconds
        
        self.init_ui()
        self.load_draft_data()
        
    def init_ui(self):
        """Initialize the user interface"""
        self.setWindowTitle("Magnus Client Intake Form v2.2")
        self.setGeometry(100, 100, 900, 700)
        
        # Set application icon
        if os.path.exists("ICON.ico"):
            self.setWindowIcon(QIcon("ICON.ico"))
        
        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        # Header
        header_layout = QHBoxLayout()
        
        # Title
        title_label = QLabel("Magnus Client Intake Form")
        title_font = QFont()
        title_font.setPointSize(18)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setStyleSheet("color: #2c3e50; margin-bottom: 10px;")
        header_layout.addWidget(title_label)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setMaximum(12)
        self.progress_bar.setValue(1)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                text-align: center;
                font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #3498db;
                border-radius: 3px;
            }
        """)
        header_layout.addWidget(self.progress_bar)
        
        main_layout.addLayout(header_layout)
        
        # Create stacked widget for pages
        self.stacked_widget = QStackedWidget()
        main_layout.addWidget(self.stacked_widget)
        
        # Create all pages
        self.create_welcome_page()           # Page 0
        self.create_personal_info_page()     # Page 1
        self.create_contact_info_page()      # Page 2
        self.create_employment_info_page()   # Page 3
        self.create_financial_info_page()    # Page 4
        self.create_spouse_info_page()       # Page 5
        self.create_dependents_page()        # Page 6
        self.create_beneficiaries_page()     # Page 7
        self.create_assets_investment_page() # Page 8
        self.create_trusted_contact_page()   # Page 9
        self.create_regulatory_page()        # Page 10
        self.create_review_submit_page()     # Page 11
        
        # Status bar
        self.statusBar().showMessage("Ready - Page 1 of 12")
        
    def create_welcome_page(self):
        """Create the welcome page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(20)
        
        # Welcome message
        welcome_label = QLabel("Welcome to Magnus Client Intake Form")
        welcome_font = QFont()
        welcome_font.setPointSize(16)
        welcome_font.setBold(True)
        welcome_label.setFont(welcome_font)
        welcome_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        welcome_label.setStyleSheet("color: #2c3e50; margin: 20px;")
        layout.addWidget(welcome_label)
        
        # Instructions
        instructions = QLabel("""
        This form will collect comprehensive information for your financial services account.
        
        Please ensure you have the following information ready:
        • Personal identification details
        • Employment and income information
        • Investment experience and objectives
        • Beneficiary information
        • Contact details for trusted persons
        
        The form includes 12 sections and takes approximately 15-20 minutes to complete.
        Your progress is automatically saved every 30 seconds.
        """)
        instructions.setWordWrap(True)
        instructions.setStyleSheet("""
            QLabel {
                background-color: #ecf0f1;
                padding: 20px;
                border-radius: 8px;
                font-size: 12px;
                line-height: 1.5;
            }
        """)
        layout.addWidget(instructions)
        
        # Navigation buttons
        layout.addLayout(self.create_navigation_buttons(back_index=None, next_index=1))
        
        self.stacked_widget.addWidget(widget)
        
    def create_personal_info_page(self):
        """Create the personal information page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Personal Information")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Form fields
        layout.addWidget(QLabel("Full Legal Name:"))
        full_name_input = EnhancedLineEdit("full_name")
        full_name_input.setObjectName("full_name")
        layout.addWidget(full_name_input)
        
        layout.addWidget(QLabel("Date of Birth:"))
        dob_input = QDateEdit()
        dob_input.setObjectName("dob")
        dob_input.setDate(QDate.currentDate().addYears(-30))
        dob_input.setCalendarPopup(True)
        dob_input.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(dob_input)
        
        layout.addWidget(QLabel("Social Security Number:"))
        ssn_input = EnhancedLineEdit("ssn")
        ssn_input.setObjectName("ssn")
        ssn_input.setPlaceholderText("XXX-XX-XXXX")
        layout.addWidget(ssn_input)
        
        layout.addWidget(QLabel("Citizenship Status:"))
        citizenship_combo = QComboBox()
        citizenship_combo.setObjectName("citizenship")
        citizenship_combo.addItems(["", "US Citizen", "Permanent Resident", "Non-Resident Alien", "Other"])
        citizenship_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(citizenship_combo)
        
        layout.addWidget(QLabel("Marital Status:"))
        marital_combo = QComboBox()
        marital_combo.setObjectName("marital_status")
        marital_combo.addItems(["", "Single", "Married", "Divorced", "Widowed", "Separated"])
        marital_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(marital_combo)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=0, next_index=2))
        
        self.stacked_widget.addWidget(widget)
        
    def create_contact_info_page(self):
        """Create the contact information page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Contact Information")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Residential Address
        layout.addWidget(QLabel("Residential Address:"))
        address_input = QTextEdit()
        address_input.setObjectName("residential_address")
        address_input.setMaximumHeight(80)
        address_input.setPlaceholderText("Street Address\nCity, State ZIP Code")
        address_input.setStyleSheet("""
            QTextEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(address_input)
        
        # Email
        layout.addWidget(QLabel("Email Address:"))
        email_input = EnhancedLineEdit("email")
        email_input.setObjectName("email")
        email_input.setPlaceholderText("example@email.com")
        layout.addWidget(email_input)
        
        # Phone numbers
        layout.addWidget(QLabel("Home Phone:"))
        home_phone_input = EnhancedLineEdit("home_phone")
        home_phone_input.setObjectName("home_phone")
        home_phone_input.setPlaceholderText("(XXX) XXX-XXXX")
        layout.addWidget(home_phone_input)
        
        layout.addWidget(QLabel("Mobile Phone:"))
        mobile_phone_input = EnhancedLineEdit("mobile_phone")
        mobile_phone_input.setObjectName("mobile_phone")
        mobile_phone_input.setPlaceholderText("(XXX) XXX-XXXX")
        layout.addWidget(mobile_phone_input)
        
        layout.addWidget(QLabel("Work Phone:"))
        work_phone_input = EnhancedLineEdit("work_phone")
        work_phone_input.setObjectName("work_phone")
        layout.addWidget(work_phone_input)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=1, next_index=3))
        
        self.stacked_widget.addWidget(widget)
        
    def create_employment_info_page(self):
        """Create the employment information page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Employment Information")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Employment Status
        layout.addWidget(QLabel("Employment Status:"))
        employment_combo = QComboBox()
        employment_combo.setObjectName("employment_status")
        employment_combo.addItems([
            "", "Employed", "Self-Employed", "Unemployed", "Retired", 
            "Student", "Homemaker", "Disabled"
        ])
        employment_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        employment_combo.currentTextChanged.connect(self.on_employment_status_changed)
        layout.addWidget(employment_combo)
        
        # Employer Information
        layout.addWidget(QLabel("Employer Name:"))
        employer_input = EnhancedLineEdit("employer_name")
        employer_input.setObjectName("employer_name")
        layout.addWidget(employer_input)
        
        layout.addWidget(QLabel("Occupation/Job Title:"))
        occupation_input = EnhancedLineEdit("occupation")
        occupation_input.setObjectName("occupation")
        layout.addWidget(occupation_input)
        
        layout.addWidget(QLabel("Years with Current Employer:"))
        years_employed_input = QSpinBox()
        years_employed_input.setObjectName("years_employed")
        years_employed_input.setRange(0, 50)
        years_employed_input.setStyleSheet("""
            QSpinBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(years_employed_input)
        
        # Annual Income
        layout.addWidget(QLabel("Annual Income:"))
        annual_income_input = EnhancedLineEdit("annual_income")
        annual_income_input.setObjectName("annual_income")
        annual_income_input.setPlaceholderText("Enter annual income in USD")
        layout.addWidget(annual_income_input)
        
        # Retirement-specific fields (initially hidden)
        self.retirement_group = QGroupBox("Retirement Information")
        self.retirement_group.setObjectName("retirement_group")
        self.retirement_group.setVisible(False)
        retirement_layout = QVBoxLayout(self.retirement_group)
        
        retirement_layout.addWidget(QLabel("Former Employer:"))
        former_employer_input = EnhancedLineEdit("former_employer")
        former_employer_input.setObjectName("former_employer")
        retirement_layout.addWidget(former_employer_input)
        
        retirement_layout.addWidget(QLabel("Source of Income:"))
        income_source_input = EnhancedLineEdit("income_source")
        income_source_input.setObjectName("income_source")
        retirement_layout.addWidget(income_source_input)
        
        layout.addWidget(self.retirement_group)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=2, next_index=4))
        
        self.stacked_widget.addWidget(widget)
        
    def on_employment_status_changed(self, text):
        """Handle employment status change to show/hide retirement fields"""
        if text == "Retired":
            self.retirement_group.setVisible(True)
        else:
            self.retirement_group.setVisible(False)
            
    def create_financial_info_page(self):
        """Create the financial information page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Financial Information")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Education Status
        layout.addWidget(QLabel("Education Status:"))
        education_combo = QComboBox()
        education_combo.setObjectName("education_status")
        education_combo.addItems([
            "", "High School", "Some College", "Associate Degree", 
            "Bachelor's Degree", "Master's Degree", "Doctorate", "Other"
        ])
        education_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(education_combo)
        
        # Tax Bracket
        layout.addWidget(QLabel("Estimated Tax Bracket:"))
        tax_bracket_combo = QComboBox()
        tax_bracket_combo.setObjectName("tax_bracket")
        tax_bracket_combo.addItems([
            "", "0-15%", "15%-32%", "32%+"
        ])
        tax_bracket_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(tax_bracket_combo)
        
        # Risk Tolerance
        layout.addWidget(QLabel("Investment Risk Tolerance:"))
        risk_combo = QComboBox()
        risk_combo.setObjectName("risk_tolerance")
        risk_combo.addItems([
            "", "Conservative", "Moderate", "Moderately Aggressive", "Aggressive"
        ])
        risk_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(risk_combo)
        
        # Investment Purpose (Updated with proper checkbox group)
        layout.addWidget(QLabel("Investment Purpose:"))
        purpose_group = QGroupBox()
        purpose_group.setStyleSheet("""
            QGroupBox {
                border: 1px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding: 10px;
            }
            QCheckBox {
                spacing: 5px;
            }
            QCheckBox::indicator {
                width: 15px;
                height: 15px;
            }
        """)
        purpose_layout = QVBoxLayout(purpose_group)
        
        purpose_options = ["Income", "Growth and Income", "Capital Appreciation", "Speculation"]
        self.purpose_checkboxes = {}
        
        for purpose in purpose_options:
            checkbox = QCheckBox(purpose)
            checkbox.setObjectName(f"investment_purpose_{purpose.lower().replace(' ', '_')}")
            self.purpose_checkboxes[purpose] = checkbox
            purpose_layout.addWidget(checkbox)
        
        layout.addWidget(purpose_group)
        
        # Investment Objectives Ranking (Updated with proper spinboxes)
        layout.addWidget(QLabel("Investment Objectives (Rank 1-5, where 1 is highest priority):"))
        objectives_group = QGroupBox()
        objectives_group.setStyleSheet("""
            QGroupBox {
                border: 1px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding: 10px;
            }
            QSpinBox {
                padding: 5px;
                border: 1px solid #bdc3c7;
                border-radius: 3px;
                min-width: 60px;
            }
        """)
        objectives_layout = QVBoxLayout(objectives_group)
        
        objectives = [
            "Trading Profits", "Speculation", "Capital Appreciation", 
            "Income", "Preservation of Capital"
        ]
        self.objective_spinboxes = {}
        
        for objective in objectives:
            h_layout = QHBoxLayout()
            label = QLabel(objective)
            label.setMinimumWidth(150)  # Ensure consistent label width
            spinbox = QSpinBox()
            spinbox.setObjectName(f"investment_objective_{objective.lower().replace(' ', '_')}")
            spinbox.setRange(1, 5)
            spinbox.setValue(3)  # Default to middle priority
            spinbox.setStyleSheet("""
                QSpinBox {
                    padding: 5px;
                    border: 1px solid #bdc3c7;
                    border-radius: 3px;
                }
            """)
            
            self.objective_spinboxes[objective] = spinbox
            h_layout.addWidget(label)
            h_layout.addWidget(spinbox)
            h_layout.addStretch()
            objectives_layout.addLayout(h_layout)
        
        layout.addWidget(objectives_group)
        
        # Net Worth
        layout.addWidget(QLabel("Estimated Net Worth (excluding primary residence):"))
        net_worth_input = EnhancedLineEdit("net_worth")
        net_worth_input.setObjectName("net_worth")
        net_worth_input.setPlaceholderText("Enter estimated net worth in USD")
        layout.addWidget(net_worth_input)
        
        # Liquid Net Worth
        layout.addWidget(QLabel("Estimated Liquid Net Worth (cash + marketable securities):"))
        liquid_net_worth_input = EnhancedLineEdit("liquid_net_worth")
        liquid_net_worth_input.setObjectName("liquid_net_worth")
        liquid_net_worth_input.setPlaceholderText("Enter estimated liquid net worth in USD")
        layout.addWidget(liquid_net_worth_input)
        
        # Assets Held Away
        layout.addWidget(QLabel("Assets Held Away (e.g., Brokerage Accounts, 401k, etc.):"))
        assets_held_away_input = EnhancedLineEdit("assets_held_away")
        assets_held_away_input.setObjectName("assets_held_away")
        assets_held_away_input.setPlaceholderText("Enter total value of assets held away in USD")
        layout.addWidget(assets_held_away_input)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=3, next_index=5))
        
        self.stacked_widget.addWidget(widget)
        
    def create_spouse_info_page(self):
        """Create the spouse information page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Spouse/Partner Information")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Checkbox for spouse applicability
        spouse_applicable_checkbox = QCheckBox("N/A (I do not have a spouse/partner)")
        spouse_applicable_checkbox.setObjectName("spouse_applicable")
        spouse_applicable_checkbox.stateChanged.connect(self.on_spouse_applicable_changed)
        layout.addWidget(spouse_applicable_checkbox)
        
        # Spouse Name
        layout.addWidget(QLabel("Full Legal Name:"))
        spouse_name_input = EnhancedLineEdit("spouse_full_name")
        spouse_name_input.setObjectName("spouse_full_name")
        layout.addWidget(spouse_name_input)
        
        # Spouse Date of Birth
        layout.addWidget(QLabel("Date of Birth:"))
        spouse_dob_input = QDateEdit()
        spouse_dob_input.setObjectName("spouse_dob")
        spouse_dob_input.setDate(QDate.currentDate().addYears(-30))
        spouse_dob_input.setCalendarPopup(True)
        spouse_dob_input.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(spouse_dob_input)
        
        # Spouse SSN
        layout.addWidget(QLabel("Social Security Number:"))
        spouse_ssn_input = EnhancedLineEdit("spouse_ssn")
        spouse_ssn_input.setObjectName("spouse_ssn")
        spouse_ssn_input.setPlaceholderText("XXX-XX-XXXX")
        layout.addWidget(spouse_ssn_input)
        
        # Spouse Employment Status
        layout.addWidget(QLabel("Employment Status:"))
        spouse_employment_combo = QComboBox()
        spouse_employment_combo.setObjectName("spouse_employment_status")
        spouse_employment_combo.addItems([
            "", "Employed", "Self-Employed", "Unemployed", "Retired", 
            "Student", "Homemaker", "Disabled"
        ])
        spouse_employment_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        layout.addWidget(spouse_employment_combo)
        
        # Spouse Employer Information
        layout.addWidget(QLabel("Employer Name:"))
        spouse_employer_input = EnhancedLineEdit("spouse_employer_name")
        spouse_employer_input.setObjectName("spouse_employer_name")
        layout.addWidget(spouse_employer_input)
        
        layout.addWidget(QLabel("Occupation/Job Title:"))
        spouse_occupation_input = EnhancedLineEdit("spouse_occupation")
        spouse_occupation_input.setObjectName("spouse_occupation")
        layout.addWidget(spouse_occupation_input)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=4, next_index=6))
        
        self.stacked_widget.addWidget(widget)
        
    def on_spouse_applicable_changed(self, state):
        """Handle spouse applicable checkbox change"""
        is_checked = state == Qt.CheckState.Checked.value
        
        # Disable/enable spouse-related fields
        for widget_name in [
            "spouse_full_name", "spouse_dob", "spouse_ssn", 
            "spouse_employment_status", "spouse_employer_name", "spouse_occupation"
        ]:
            widget = self.findChild(QWidget, widget_name)
            if widget:
                widget.setEnabled(not is_checked)
                
    def create_dependents_page(self):
        """Create the dependents information page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Dependents Information")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Dependents list container
        self.dependents_layout = QVBoxLayout()
        self.dependents_layout.setSpacing(10)
        
        self.dependents_scroll_area = QScrollArea()
        self.dependents_scroll_area.setWidgetResizable(True)
        self.dependents_scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        
        dependents_container = QWidget()
        dependents_container.setLayout(self.dependents_layout)
        self.dependents_scroll_area.setWidget(dependents_container)
        
        layout.addWidget(self.dependents_scroll_area)
        
        # Add Dependent button
        add_dependent_btn = QPushButton("Add Dependent")
        add_dependent_btn.setStyleSheet("""
            QPushButton {
                background-color: #17a2b8;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #138496;
            }
        """)
        add_dependent_btn.clicked.connect(self.add_dependent_field)
        layout.addWidget(add_dependent_btn)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=5, next_index=7))
        
        self.stacked_widget.addWidget(widget)
        
    def add_dependent_field(self, dependent_data=None):
        """Add fields for a new dependent"""
        dependent_frame = QFrame()
        dependent_frame.setFrameShape(QFrame.Shape.StyledPanel)
        dependent_frame.setFrameShadow(QFrame.Shadow.Raised)
        dependent_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                background-color: #f0f0f0;
            }
        """)
        
        frame_layout = QVBoxLayout(dependent_frame)
        
        # Name
        frame_layout.addWidget(QLabel("Dependent Full Name:"))
        name_input = EnhancedLineEdit("dependent_name")
        name_input.setObjectName(f"dependent_name_{self.dependents_layout.count()}")
        frame_layout.addWidget(name_input)
        
        # Date of Birth
        frame_layout.addWidget(QLabel("Dependent Date of Birth:"))
        dob_input = QDateEdit()
        dob_input.setObjectName(f"dependent_dob_{self.dependents_layout.count()}")
        dob_input.setDate(QDate.currentDate().addYears(-10))
        dob_input.setCalendarPopup(True)
        dob_input.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        frame_layout.addWidget(dob_input)
        
        # Relationship
        frame_layout.addWidget(QLabel("Relationship:"))
        relationship_input = EnhancedLineEdit("dependent_relationship")
        relationship_input.setObjectName(f"dependent_relationship_{self.dependents_layout.count()}")
        frame_layout.addWidget(relationship_input)
        
        # Remove button
        remove_btn = QPushButton("Remove Dependent")
        remove_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        remove_btn.clicked.connect(lambda: self.remove_dependent_field(dependent_frame))
        frame_layout.addWidget(remove_btn)
        
        self.dependents_layout.addWidget(dependent_frame)
        
        if dependent_data:
            name_input.setText(dependent_data.get("name", ""))
            dob_input.setDate(QDate.fromString(dependent_data.get("dob", ""), "MM/dd/yyyy"))
            relationship_input.setText(dependent_data.get("relationship", ""))
            
    def remove_dependent_field(self, frame):
        """Remove dependent fields"""
        frame.deleteLater()
        
    def create_beneficiaries_page(self):
        """Create the beneficiaries information page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Beneficiaries Information")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Beneficiaries list container
        self.beneficiaries_layout = QVBoxLayout()
        self.beneficiaries_layout.setSpacing(10)
        
        self.beneficiaries_scroll_area = QScrollArea()
        self.beneficiaries_scroll_area.setWidgetResizable(True)
        self.beneficiaries_scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        
        beneficiaries_container = QWidget()
        beneficiaries_container.setLayout(self.beneficiaries_layout)
        self.beneficiaries_scroll_area.setWidget(beneficiaries_container)
        
        layout.addWidget(self.beneficiaries_scroll_area)
        
        # Add Beneficiary button
        add_beneficiary_btn = QPushButton("Add Beneficiary")
        add_beneficiary_btn.setStyleSheet("""
            QPushButton {
                background-color: #17a2b8;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #138496;
            }
        """)
        add_beneficiary_btn.clicked.connect(self.add_beneficiary_field)
        layout.addWidget(add_beneficiary_btn)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=6, next_index=8))
        
        self.stacked_widget.addWidget(widget)
        
    def add_beneficiary_field(self, beneficiary_data=None):
        """Add fields for a new beneficiary"""
        beneficiary_frame = QFrame()
        beneficiary_frame.setFrameShape(QFrame.Shape.StyledPanel)
        beneficiary_frame.setFrameShadow(QFrame.Shadow.Raised)
        beneficiary_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                background-color: #f0f0f0;
            }
        """)
        
        frame_layout = QVBoxLayout(beneficiary_frame)
        
        # Name
        frame_layout.addWidget(QLabel("Beneficiary Full Name:"))
        name_input = EnhancedLineEdit("beneficiary_name")
        name_input.setObjectName(f"beneficiary_name_{self.beneficiaries_layout.count()}")
        frame_layout.addWidget(name_input)
        
        # Date of Birth
        frame_layout.addWidget(QLabel("Beneficiary Date of Birth:"))
        dob_input = QDateEdit()
        dob_input.setObjectName(f"beneficiary_dob_{self.beneficiaries_layout.count()}")
        dob_input.setDate(QDate.currentDate().addYears(-10))
        dob_input.setCalendarPopup(True)
        dob_input.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        frame_layout.addWidget(dob_input)
        
        # Relationship
        frame_layout.addWidget(QLabel("Relationship:"))
        relationship_input = EnhancedLineEdit("beneficiary_relationship")
        relationship_input.setObjectName(f"beneficiary_relationship_{self.beneficiaries_layout.count()}")
        frame_layout.addWidget(relationship_input)
        
        # Percentage
        frame_layout.addWidget(QLabel("Allocation Percentage (%):"))
        percentage_spin = QSpinBox()
        percentage_spin.setObjectName(f"beneficiary_percentage_{self.beneficiaries_layout.count()}")
        percentage_spin.setRange(0, 100)
        percentage_spin.setSuffix("%")
        percentage_spin.setStyleSheet("""
            QSpinBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        frame_layout.addWidget(percentage_spin)
        
        # Remove button
        remove_btn = QPushButton("Remove Beneficiary")
        remove_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        remove_btn.clicked.connect(lambda: self.remove_beneficiary_field(beneficiary_frame))
        frame_layout.addWidget(remove_btn)
        
        self.beneficiaries_layout.addWidget(beneficiary_frame)
        
        if beneficiary_data:
            name_input.setText(beneficiary_data.get("name", ""))
            dob_input.setDate(QDate.fromString(beneficiary_data.get("dob", ""), "MM/dd/yyyy"))
            relationship_input.setText(beneficiary_data.get("relationship", ""))
            percentage_spin.setValue(beneficiary_data.get("percentage", 0))
            
    def remove_beneficiary_field(self, frame):
        """Remove beneficiary fields"""
        frame.deleteLater()
        
    def create_assets_investment_page(self):
        """Create the assets and investment experience page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Assets & Investment Experience")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Create a scroll area for the entire content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        
        # Create a widget to hold all the content
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        
        # Asset Breakdown (optional)
        include_breakdown_checkbox = QCheckBox("Include Asset Breakdown")
        include_breakdown_checkbox.setObjectName("include_breakdown")
        include_breakdown_checkbox.stateChanged.connect(self.on_include_breakdown_changed)
        content_layout.addWidget(include_breakdown_checkbox)
        
        self.asset_breakdown_group = QGroupBox("Asset Breakdown")
        self.asset_breakdown_group.setObjectName("asset_breakdown_group")
        self.asset_breakdown_group.setVisible(False)
        
        breakdown_layout = QVBoxLayout(self.asset_breakdown_group)
        self.asset_breakdown_fields = {}
        
        asset_types = [
            "Stocks", "Bonds", "Mutual Funds", "ETFs", "UITs", 
            "Annuities (Fixed)", "Annuities (Variable)", "Options", 
            "Commodities", "Alternative Investments", "Limited Partnerships", 
            "Variable Contracts", "Short-Term", "Other"
        ]
        
        for asset_type in asset_types:
            h_layout = QHBoxLayout()
            label = QLabel(f"{asset_type} (%):")
            spin_box = QSpinBox()
            spin_box.setObjectName(f"asset_breakdown_{asset_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}")
            spin_box.setRange(0, 100)
            spin_box.setSuffix("%")
            
            self.asset_breakdown_fields[asset_type] = spin_box
            h_layout.addWidget(label)
            h_layout.addWidget(spin_box)
            breakdown_layout.addLayout(h_layout)
            
        content_layout.addWidget(self.asset_breakdown_group)
        
        # Investment Experience by Asset Type
        experience_label = QLabel("Investment Experience by Asset Type:")
        experience_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        content_layout.addWidget(experience_label)
        
        # Create a scroll area specifically for investment experience
        experience_scroll = QScrollArea()
        experience_scroll.setWidgetResizable(True)
        experience_scroll.setFrameShape(QFrame.Shape.NoFrame)
        experience_scroll.setMinimumHeight(300)  # Set minimum height for better visibility
        
        experience_widget = QWidget()
        self.asset_experience_layout = QVBoxLayout(experience_widget)
        self.asset_experience_fields = {}
        
        experience_types = [
            "Stocks", "Bonds", "Mutual Funds", "UITs", 
            "Annuities (Fixed)", "Annuities (Variable)", "Options", 
            "Commodities", "Alternative Investments", "Limited Partnerships", 
            "Variable Contracts"
        ]
        
        for exp_type in experience_types:
            group_box = QGroupBox(exp_type)
            group_box.setStyleSheet("""
                QGroupBox {
                    font-weight: bold;
                    border: 1px solid #bdc3c7;
                    border-radius: 5px;
                    margin-top: 10px;
                    padding-top: 10px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 5px 0 5px;
                }
            """)
            group_box_layout = QHBoxLayout(group_box)
            
            year_label = QLabel("Year Started:")
            year_input = QLineEdit()
            year_input.setObjectName(f"asset_experience_{exp_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}_year")
            year_input.setPlaceholderText("YYYY")
            year_input.setMaximumWidth(80)
            year_input.setStyleSheet("""
                QLineEdit {
                    padding: 5px;
                    border: 1px solid #bdc3c7;
                    border-radius: 3px;
                }
            """)
            
            level_label = QLabel("Level:")
            level_combo = QComboBox()
            level_combo.setObjectName(f"asset_experience_{exp_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}_level")
            level_combo.addItems(["", "None", "Limited", "Good", "Extensive"])
            level_combo.setStyleSheet("""
                QComboBox {
                    padding: 5px;
                    border: 1px solid #bdc3c7;
                    border-radius: 3px;
                }
            """)
            
            group_box_layout.addWidget(year_label)
            group_box_layout.addWidget(year_input)
            group_box_layout.addWidget(level_label)
            group_box_layout.addWidget(level_combo)
            group_box_layout.addStretch()
            
            self.asset_experience_fields[exp_type] = {"year": year_input, "level": level_combo}
            self.asset_experience_layout.addWidget(group_box)
        
        experience_scroll.setWidget(experience_widget)
        content_layout.addWidget(experience_scroll)
        
        # Outside Broker Firm
        has_outside_broker_checkbox = QCheckBox("Do you have assets with an outside broker firm?")
        has_outside_broker_checkbox.setObjectName("has_outside_broker")
        has_outside_broker_checkbox.stateChanged.connect(self.on_has_outside_broker_changed)
        content_layout.addWidget(has_outside_broker_checkbox)
        
        self.outside_broker_group = QGroupBox("Outside Broker Firm Information")
        self.outside_broker_group.setObjectName("outside_broker_group")
        self.outside_broker_group.setVisible(False)
        
        outside_broker_layout = QVBoxLayout(self.outside_broker_group)
        
        outside_broker_layout.addWidget(QLabel("Firm Name:"))
        outside_firm_name_input = EnhancedLineEdit("outside_firm_name")
        outside_firm_name_input.setObjectName("outside_firm_name")
        outside_broker_layout.addWidget(outside_firm_name_input)
        
        outside_broker_layout.addWidget(QLabel("Account Type:"))
        outside_account_type_input = QComboBox()
        outside_account_type_input.setObjectName("outside_broker_account_type")
        outside_account_type_input.addItems([
            "", "Individual", "Joint", "IRA", "Roth IRA", 
            "401(k)", "Trust", "Other"
        ])
        outside_broker_layout.addWidget(outside_account_type_input)
        
        outside_broker_layout.addWidget(QLabel("Account Number:"))
        outside_account_number_input = EnhancedLineEdit("outside_broker_account_number")
        outside_account_number_input.setObjectName("outside_broker_account_number")
        outside_broker_layout.addWidget(outside_account_number_input)
        
        outside_broker_layout.addWidget(QLabel("Liquid Amount with this Firm:"))
        outside_liquid_amount_input = EnhancedLineEdit("outside_liquid_amount")
        outside_liquid_amount_input.setObjectName("outside_liquid_amount")
        outside_liquid_amount_input.setPlaceholderText("Enter liquid amount in USD")
        outside_broker_layout.addWidget(outside_liquid_amount_input)
        
        content_layout.addWidget(self.outside_broker_group)
        
        # Add the content widget to the scroll area
        scroll_area.setWidget(content_widget)
        layout.addWidget(scroll_area)
        
        # Navigation buttons
        layout.addLayout(self.create_navigation_buttons(back_index=7, next_index=9))
        
        self.stacked_widget.addWidget(widget)
        
    def on_include_breakdown_changed(self, state):
        """Handle include breakdown checkbox change"""
        is_checked = state == Qt.CheckState.Checked.value
        self.asset_breakdown_group.setVisible(is_checked)
        
    def on_has_outside_broker_changed(self, state):
        """Handle has outside broker checkbox change"""
        is_checked = state == Qt.CheckState.Checked.value
        self.outside_broker_group.setVisible(is_checked)
        
    def create_trusted_contact_page(self):
        """Create the trusted contact person page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Trusted Contact Person")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Instructions
        instructions = QLabel("""
        Please provide information for a trusted contact person. This person may be contacted 
        in the event we are unable to reach you, or if we have concerns about your health 
        or financial exploitation.
        """)
        instructions.setWordWrap(True)
        instructions.setStyleSheet("font-style: italic; color: #7f8c8d; margin-bottom: 15px;")
        layout.addWidget(instructions)
        
        # Trusted Contact Name
        layout.addWidget(QLabel("Full Legal Name:"))
        trusted_name_input = EnhancedLineEdit("trusted_full_name")
        trusted_name_input.setObjectName("trusted_full_name")
        layout.addWidget(trusted_name_input)
        
        # Trusted Contact Relationship
        layout.addWidget(QLabel("Relationship to You:"))
        trusted_relationship_input = EnhancedLineEdit("trusted_relationship")
        trusted_relationship_input.setObjectName("trusted_relationship")
        layout.addWidget(trusted_relationship_input)
        
        # Trusted Contact Phone
        layout.addWidget(QLabel("Phone Number:"))
        trusted_phone_input = EnhancedLineEdit("trusted_phone")
        trusted_phone_input.setObjectName("trusted_phone")
        trusted_phone_input.setPlaceholderText("(XXX) XXX-XXXX")
        layout.addWidget(trusted_phone_input)
        
        # Trusted Contact Email
        layout.addWidget(QLabel("Email Address:"))
        trusted_email_input = EnhancedLineEdit("trusted_email")
        trusted_email_input.setObjectName("trusted_email")
        trusted_email_input.setPlaceholderText("example@email.com")
        layout.addWidget(trusted_email_input)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=8, next_index=10))
        
        self.stacked_widget.addWidget(widget)
        
    def create_regulatory_page(self):
        """Create the regulatory consent page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Regulatory Consent")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Electronic Delivery Consent
        layout.addWidget(QLabel("Electronic Delivery Consent:"))
        
        reg_group = QButtonGroup()
        
        reg_yes = QRadioButton("Yes - I consent to receive regulatory communications electronically")
        reg_yes.setObjectName("electronic_regulatory_yes")
        reg_group.addButton(reg_yes)
        layout.addWidget(reg_yes)
        
        reg_no = QRadioButton("No - I prefer to receive regulatory communications by mail")
        reg_no.setObjectName("electronic_regulatory_no")
        reg_group.addButton(reg_no)
        layout.addWidget(reg_no)
        
        # Disclosure text
        disclosure = QLabel("""
        Electronic Delivery Disclosure:
        
        By selecting "Yes" above, you consent to receive regulatory communications, 
        account statements, confirmations, prospectuses, and other important documents 
        electronically. You may withdraw this consent at any time by contacting us.
        
        Electronic delivery helps reduce paper waste and provides faster access to 
        your important documents. You will receive email notifications when new 
        documents are available in your secure online account.
        
        System Requirements: You must have access to a computer with internet 
        connection and email capability to receive electronic communications.
        """)
        disclosure.setWordWrap(True)
        disclosure.setStyleSheet("""
            QLabel {
                background-color: #f8f9fa;
                padding: 15px;
                border-radius: 5px;
                font-size: 11px;
                line-height: 1.4;
                border: 1px solid #dee2e6;
            }
        """)
        layout.addWidget(disclosure)
        
        layout.addStretch()
        layout.addLayout(self.create_navigation_buttons(back_index=9, next_index=11))
        
        self.stacked_widget.addWidget(widget)
        
    def create_review_submit_page(self):
        """Create the review and submit page"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Page title
        title = QLabel("Review & Submit")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #2c3e50; margin-bottom: 15px;")
        layout.addWidget(title)
        
        # Instructions
        instructions = QLabel("Please review your information and submit the form to generate your PDF report.")
        instructions.setStyleSheet("font-style: italic; color: #7f8c8d; margin-bottom: 15px;")
        layout.addWidget(instructions)
        
        # Review area
        self.review_area = QTextEdit()
        self.review_area.setReadOnly(True)
        self.review_area.setStyleSheet("""
            QTextEdit {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 5px;
                padding: 10px;
                font-family: monospace;
                font-size: 11px;
            }
        """)
        layout.addWidget(self.review_area)
        
        # Action buttons
        button_layout = QHBoxLayout()
        
        save_draft_btn = QPushButton("Save Draft")
        save_draft_btn.clicked.connect(self.save_draft)
        save_draft_btn.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
        """)
        button_layout.addWidget(save_draft_btn)
        
        generate_pdf_btn = QPushButton("Generate PDF Report")
        generate_pdf_btn.clicked.connect(self.generate_pdf_report)
        generate_pdf_btn.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
        """)
        button_layout.addWidget(generate_pdf_btn)
        
        layout.addLayout(button_layout)
        layout.addLayout(self.create_navigation_buttons(back_index=10, next_index=None))
        
        self.stacked_widget.addWidget(widget)
        
    def create_navigation_buttons(self, back_index=None, next_index=None):
        """Create navigation buttons layout"""
        layout = QHBoxLayout()
        
        if back_index is not None:
            back_btn = QPushButton("← Back")
            back_btn.clicked.connect(lambda: self.navigate_to_page(back_index))
            back_btn.setStyleSheet("""
                QPushButton {
                    background-color: #6c757d;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 5px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #5a6268;
                }
            """)
            layout.addWidget(back_btn)
        
        layout.addStretch()
        
        if next_index is not None:
            next_btn = QPushButton("Next →")
            next_btn.clicked.connect(lambda: self.navigate_to_page(next_index))
            next_btn.setStyleSheet("""
                QPushButton {
                    background-color: #007bff;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 5px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #0056b3;
                }
            """)
            layout.addWidget(next_btn)
        
        return layout
        
    def navigate_to_page(self, page_index):
        """Navigate to a specific page"""
        if page_index == 11:  # Review page
            self.update_review_area()
        
        self.current_page = page_index
        self.stacked_widget.setCurrentIndex(page_index)
        self.progress_bar.setValue(page_index + 1)
        self.statusBar().showMessage(f"Page {page_index + 1} of 12")
        
        # Auto-save when navigating
        self.collect_form_data()
        
    def update_review_area(self):
        """Update the review area with current form data"""
        self.collect_form_data()
        
        review_text = "=== MAGNUS CLIENT INTAKE FORM - REVIEW ===\n\n"
        
        # Helper to format fields
        def format_field(label, value):
            return f"  {label}: {value if value else '[Not provided]'}\n"

        # Personal Information
        review_text += "PERSONAL INFORMATION:\n"
        review_text += format_field("Full Name", self.form_data.get("full_name"))
        review_text += format_field("Date of Birth", self.form_data.get("dob"))
        review_text += format_field("Social Security Number", self.form_data.get("ssn"))
        review_text += format_field("Citizenship", self.form_data.get("citizenship"))
        review_text += format_field("Marital Status", self.form_data.get("marital_status"))
        review_text += "\n"

        # Contact Information
        review_text += "CONTACT INFORMATION:\n"
        review_text += format_field("Residential Address", self.form_data.get("residential_address"))
        if self.form_data.get("mailing_address_different"):
            review_text += format_field("Mailing Address", self.form_data.get("mailing_address"))
        review_text += format_field("Email", self.form_data.get("email"))
        review_text += format_field("Home Phone", self.form_data.get("home_phone"))
        review_text += format_field("Mobile Phone", self.form_data.get("mobile_phone"))
        review_text += format_field("Work Phone", self.form_data.get("work_phone"))
        review_text += "\n"

        # Employment Information
        review_text += "EMPLOYMENT INFORMATION:\n"
        review_text += format_field("Employment Status", self.form_data.get("employment_status"))
        review_text += format_field("Employer Name", self.form_data.get("employer_name"))
        review_text += format_field("Occupation", self.form_data.get("occupation"))
        review_text += format_field("Years Employed", self.form_data.get("years_employed"))
        review_text += format_field("Annual Income", self.form_data.get("annual_income"))
        review_text += format_field("Employer Address", self.form_data.get("employer_address"))
        review_text += "\n"

        # Retirement Information
        if self.form_data.get("employment_status") == "Retired":
            review_text += "RETIREMENT INFORMATION:\n"
            review_text += format_field("Former Employer", self.form_data.get("former_employer"))
            review_text += format_field("Source of Income", self.form_data.get("income_source"))
            review_text += "\n"

        # Financial Information
        review_text += "FINANCIAL INFORMATION:\n"
        review_text += format_field("Education Status", self.form_data.get("education_status"))
        review_text += format_field("Estimated Tax Bracket", self.form_data.get("tax_bracket"))
        review_text += format_field("Investment Risk Tolerance", self.form_data.get("risk_tolerance"))
        review_text += format_field("Investment Purpose", self.form_data.get("investment_purpose"))
        review_text += format_field("Investment Objectives", self.form_data.get("investment_objective"))
        review_text += format_field("Net Worth (excluding primary home)", self.form_data.get("net_worth"))
        review_text += format_field("Liquid Net Worth", self.form_data.get("liquid_net_worth"))
        review_text += format_field("Assets Held Away", self.form_data.get("assets_held_away"))
        review_text += "\n"

        # Spouse Information
        if self.form_data.get("spouse_applicable"):
            review_text += "SPOUSE INFORMATION:\n"
            review_text += format_field("Spouse Full Name", self.form_data.get("spouse_full_name"))
            review_text += format_field("Spouse Date of Birth", self.form_data.get("spouse_dob"))
            review_text += format_field("Spouse SSN", self.form_data.get("spouse_ssn"))
            review_text += format_field("Spouse Employment Status", self.form_data.get("spouse_employment_status"))
            review_text += format_field("Spouse Employer Name", self.form_data.get("spouse_employer_name"))
            review_text += format_field("Spouse Occupation/Title", self.form_data.get("spouse_occupation"))
            review_text += "\n"

        # Dependents
        review_text += "DEPENDENTS:\n"
        dependents = self.form_data.get("dependents", [])
        if dependents:
            for i, dep in enumerate(dependents):
                review_text += f"  Dependent {i+1}:\n"
                review_text += format_field("    Name", dep.get("name"))
                review_text += format_field("    Date of Birth", dep.get("dob"))
                review_text += format_field("    Relationship", dep.get("relationship"))
        else:
            review_text += "  [No dependents specified]\n"
        review_text += "\n"

        # Beneficiaries
        review_text += "BENEFICIARIES:\n"
        beneficiaries = self.form_data.get("beneficiaries", [])
        if beneficiaries:
            for i, ben in enumerate(beneficiaries):
                review_text += f"  Beneficiary {i+1}:\n"
                review_text += format_field("    Name", ben.get("name"))
                review_text += format_field("    Date of Birth", ben.get("dob"))
                review_text += format_field("    Relationship", ben.get("relationship"))
                percentage = ben.get('percentage', '')
                review_text += format_field("    Percentage", f"{percentage}%" if percentage else "[Not provided]")
        else:
            review_text += "  [No beneficiaries specified]\n"
        review_text += "\n"

        # Asset Breakdown
        review_text += "ASSET BREAKDOWN:\n"
        asset_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures", "Short-Term", "Other"]
        for asset_type in asset_types:
            field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_')}"
            value = self.form_data.get(field_name)
            review_text += format_field(asset_type, f"{value}%" if value else None)
        review_text += "\n"

        # Investment Experience
        review_text += "INVESTMENT EXPERIENCE:\n"
        experience_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures"]
        for exp_type in experience_types:
            year_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_year"
            level_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_level"
            
            year = self.form_data.get(year_field)
            level = self.form_data.get(level_field)
            
            review_text += f"  {exp_type}:\n"
            review_text += format_field("    Year Started", year)
            review_text += format_field("    Experience Level", level)
        review_text += "\n"

        # Outside Broker Information
        if self.form_data.get("has_outside_broker"):
            review_text += "OUTSIDE BROKER INFORMATION:\n"
            review_text += format_field("Broker Firm Name", self.form_data.get("outside_firm_name"))
            review_text += format_field("Account Number", self.form_data.get("outside_broker_account_number"))
            review_text += format_field("Account Type", self.form_data.get("outside_broker_account_type"))
            review_text += "\n"

        # Trusted Contact Information
        review_text += "TRUSTED CONTACT INFORMATION:\n"
        review_text += format_field("Full Name", self.form_data.get("trusted_full_name"))
        review_text += format_field("Relationship", self.form_data.get("trusted_relationship"))
        review_text += format_field("Phone Number", self.form_data.get("trusted_phone"))
        review_text += format_field("Email Address", self.form_data.get("trusted_email"))
        review_text += "\n"

        # Regulatory Consent
        review_text += "REGULATORY CONSENT:\n"
        electronic_consent = "Yes" if self.form_data.get("electronic_regulatory_yes") else "No"
        review_text += format_field("Electronic Delivery Consent", electronic_consent)
        review_text += "\n"

        self.review_area.setPlainText(review_text)
        
    def collect_form_data(self):
        """Collect all form data from the UI"""
        # Get all widgets with object names
        for widget in self.findChildren((QLineEdit, QComboBox, QDateEdit, QTextEdit, QSpinBox, QCheckBox, QRadioButton)):
            object_name = widget.objectName()
            if object_name:
                if isinstance(widget, QLineEdit):
                    self.form_data[object_name] = widget.text()
                elif isinstance(widget, QComboBox):
                    self.form_data[object_name] = widget.currentText()
                elif isinstance(widget, QDateEdit):
                    self.form_data[object_name] = widget.date().toString("MM/dd/yyyy")
                elif isinstance(widget, QTextEdit):
                    self.form_data[object_name] = widget.toPlainText()
                elif isinstance(widget, QSpinBox):
                    self.form_data[object_name] = widget.value()
                elif isinstance(widget, QCheckBox):
                    self.form_data[object_name] = widget.isChecked()
                elif isinstance(widget, QRadioButton):
                    if widget.isChecked():
                        self.form_data[object_name] = True

        # Collect dependents data
        dependents = []
        for i in range(self.dependents_layout.count()):
            frame = self.dependents_layout.itemAt(i).widget()
            if isinstance(frame, QFrame):
                dependent_data = {}
                for child in frame.findChildren((QLineEdit, QDateEdit)):
                    if isinstance(child, QLineEdit):
                        if "name" in child.objectName():
                            dependent_data["name"] = child.text()
                        elif "relationship" in child.objectName():
                            dependent_data["relationship"] = child.text()
                    elif isinstance(child, QDateEdit):
                        dependent_data["dob"] = child.date().toString("MM/dd/yyyy")
                if dependent_data:
                    dependents.append(dependent_data)
        self.form_data["dependents"] = dependents

        # Collect beneficiaries data
        beneficiaries = []
        for i in range(self.beneficiaries_layout.count()):
            frame = self.beneficiaries_layout.itemAt(i).widget()
            if isinstance(frame, QFrame):
                beneficiary_data = {}
                for child in frame.findChildren((QLineEdit, QDateEdit, QSpinBox)):
                    if isinstance(child, QLineEdit):
                        if "name" in child.objectName():
                            beneficiary_data["name"] = child.text()
                        elif "relationship" in child.objectName():
                            beneficiary_data["relationship"] = child.text()
                    elif isinstance(child, QDateEdit):
                        beneficiary_data["dob"] = child.date().toString("MM/dd/yyyy")
                    elif isinstance(child, QSpinBox):
                        beneficiary_data["percentage"] = child.value()
                if beneficiary_data:
                    beneficiaries.append(beneficiary_data)
        self.form_data["beneficiaries"] = beneficiaries

        # Collect asset breakdown data
        if hasattr(self, 'asset_breakdown_fields'):
            for asset_type, spin_box in self.asset_breakdown_fields.items():
                field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_')}"
                self.form_data[field_name] = spin_box.value()

        # Collect investment experience data
        if hasattr(self, 'asset_experience_fields'):
            for exp_type, fields in self.asset_experience_fields.items():
                year_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_year"
                level_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_level"
                self.form_data[year_field] = fields["year"].text()
                self.form_data[level_field] = fields["level"].currentText()
        
    def auto_save_data(self):
        """Auto-save form data"""
        # Removed auto-save functionality to prevent JSON format saving
        pass
            
    def load_draft_data(self):
        """Load draft data if available"""
        try:
            temp_file = os.path.join(tempfile.gettempdir(), "magnus_form_autosave.json")
            if os.path.exists(temp_file):
                with open(temp_file, 'r') as f:
                    self.form_data = json.load(f)
                self.populate_form_fields()
        except Exception as e:
            print(f"Failed to load draft: {e}")
            
    def populate_form_fields(self):
        """Populate form fields with loaded data"""
        for object_name, value in self.form_data.items():
            widget = self.findChild((QLineEdit, QComboBox, QDateEdit, QTextEdit, QSpinBox, QCheckBox, QRadioButton), object_name)
            if widget:
                try:
                    if isinstance(widget, QLineEdit):
                        widget.setText(str(value))
                    elif isinstance(widget, QComboBox):
                        index = widget.findText(str(value))
                        if index >= 0:
                            widget.setCurrentIndex(index)
                    elif isinstance(widget, QDateEdit):
                        widget.setDate(QDate.fromString(str(value), "MM/dd/yyyy"))
                    elif isinstance(widget, QTextEdit):
                        widget.setPlainText(str(value))
                    elif isinstance(widget, QSpinBox):
                        if value:
                            widget.setValue(int(value))
                    elif isinstance(widget, QCheckBox):
                        widget.setChecked(bool(value))
                    elif isinstance(widget, QRadioButton):
                        widget.setChecked(bool(value))
                except Exception as e:
                    print(f"Failed to populate field {object_name}: {e}")
                    
    def save_draft(self):
        """Save current form as draft in Word format"""
        self.collect_form_data()
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Draft", "magnus_form_draft.docx", "Word Files (*.docx)"
        )
        
        if file_path:
            try:
                doc = Document()
                
                # Title
                title = doc.add_heading('Magnus Client Intake Form', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Helper function to format monetary values
                def format_money(value):
                    if value:
                        try:
                            return f"${int(value):,}"
                        except ValueError:
                            return value
                    return "[Not provided]"
                
                # Personal Information
                doc.add_heading('Personal Information', level=1)
                doc.add_paragraph(f"Full Name: {self.form_data.get('full_name', '[Not provided]')}")
                doc.add_paragraph(f"Date of Birth: {self.form_data.get('dob', '[Not provided]')}")
                doc.add_paragraph(f"Social Security Number: {self.form_data.get('ssn', '[Not provided]')}")
                doc.add_paragraph(f"Citizenship: {self.form_data.get('citizenship', '[Not provided]')}")
                doc.add_paragraph(f"Marital Status: {self.form_data.get('marital_status', '[Not provided]')}")
                doc.add_paragraph()
                
                # Contact Information
                doc.add_heading('Contact Information', level=1)
                doc.add_paragraph(f"Residential Address: {self.form_data.get('residential_address', '[Not provided]')}")
                doc.add_paragraph(f"Email: {self.form_data.get('email', '[Not provided]')}")
                doc.add_paragraph(f"Home Phone: {self.form_data.get('home_phone', '[Not provided]')}")
                doc.add_paragraph(f"Mobile Phone: {self.form_data.get('mobile_phone', '[Not provided]')}")
                doc.add_paragraph(f"Work Phone: {self.form_data.get('work_phone', '[Not provided]')}")
                doc.add_paragraph()
                
                # Employment Information
                doc.add_heading('Employment Information', level=1)
                doc.add_paragraph(f"Employment Status: {self.form_data.get('employment_status', '[Not provided]')}")
                doc.add_paragraph(f"Employer Name: {self.form_data.get('employer_name', '[Not provided]')}")
                doc.add_paragraph(f"Occupation: {self.form_data.get('occupation', '[Not provided]')}")
                doc.add_paragraph(f"Years Employed: {self.form_data.get('years_employed', '[Not provided]')}")
                doc.add_paragraph(f"Annual Income: {format_money(self.form_data.get('annual_income'))}")
                
                # Retirement Information (if applicable)
                if self.form_data.get('employment_status') == 'Retired':
                    doc.add_heading('Retirement Information', level=1)
                    doc.add_paragraph(f"Former Employer: {self.form_data.get('former_employer', '[Not provided]')}")
                    doc.add_paragraph(f"Source of Income: {self.form_data.get('income_source', '[Not provided]')}")
                doc.add_paragraph()
                
                # Financial Information
                doc.add_heading('Financial Information', level=1)
                doc.add_paragraph(f"Education Status: {self.form_data.get('education_status', '[Not provided]')}")
                doc.add_paragraph(f"Estimated Tax Bracket: {self.form_data.get('tax_bracket', '[Not provided]')}")
                doc.add_paragraph(f"Investment Risk Tolerance: {self.form_data.get('risk_tolerance', '[Not provided]')}")
                
                # Investment Purpose
                doc.add_paragraph("Investment Purpose:")
                purpose_list = []
                for purpose in ["Income", "Growth and Income", "Capital Appreciation", "Speculation"]:
                    if self.form_data.get(f"investment_purpose_{purpose.lower().replace(' ', '_')}"):
                        purpose_list.append(purpose)
                doc.add_paragraph(", ".join(purpose_list) if purpose_list else "[Not provided]")
                
                # Investment Objectives
                doc.add_paragraph("Investment Objectives (Ranked 1-5):")
                objectives = [
                    "Trading Profits", "Speculation", "Capital Appreciation", 
                    "Income", "Preservation of Capital"
                ]
                for objective in objectives:
                    rank = self.form_data.get(f"investment_objective_{objective.lower().replace(' ', '_')}")
                    if rank:
                        doc.add_paragraph(f"  {objective}: {rank}")
                
                doc.add_paragraph(f"Net Worth: {format_money(self.form_data.get('net_worth'))}")
                doc.add_paragraph(f"Liquid Net Worth: {format_money(self.form_data.get('liquid_net_worth'))}")
                doc.add_paragraph(f"Assets Held Away: {format_money(self.form_data.get('assets_held_away'))}")
                doc.add_paragraph()
                
                # Spouse Information
                if self.form_data.get('spouse_applicable'):
                    doc.add_heading('Spouse Information', level=1)
                    doc.add_paragraph(f"Full Name: {self.form_data.get('spouse_full_name', '[Not provided]')}")
                    doc.add_paragraph(f"Date of Birth: {self.form_data.get('spouse_dob', '[Not provided]')}")
                    doc.add_paragraph(f"Social Security Number: {self.form_data.get('spouse_ssn', '[Not provided]')}")
                    doc.add_paragraph(f"Employment Status: {self.form_data.get('spouse_employment_status', '[Not provided]')}")
                    doc.add_paragraph(f"Employer Name: {self.form_data.get('spouse_employer_name', '[Not provided]')}")
                    doc.add_paragraph(f"Occupation: {self.form_data.get('spouse_occupation', '[Not provided]')}")
                    doc.add_paragraph()
                
                # Dependents
                doc.add_heading('Dependents', level=1)
                dependents = self.form_data.get('dependents', [])
                if dependents:
                    for i, dep in enumerate(dependents, 1):
                        doc.add_paragraph(f"Dependent {i}:")
                        doc.add_paragraph(f"  Name: {dep.get('name', '[Not provided]')}")
                        doc.add_paragraph(f"  Date of Birth: {dep.get('dob', '[Not provided]')}")
                        doc.add_paragraph(f"  Relationship: {dep.get('relationship', '[Not provided]')}")
                else:
                    doc.add_paragraph("[No dependents specified]")
                doc.add_paragraph()
                
                # Beneficiaries
                doc.add_heading('Beneficiaries', level=1)
                beneficiaries = self.form_data.get('beneficiaries', [])
                if beneficiaries:
                    for i, ben in enumerate(beneficiaries, 1):
                        doc.add_paragraph(f"Beneficiary {i}:")
                        doc.add_paragraph(f"  Name: {ben.get('name', '[Not provided]')}")
                        doc.add_paragraph(f"  Date of Birth: {ben.get('dob', '[Not provided]')}")
                        doc.add_paragraph(f"  Relationship: {ben.get('relationship', '[Not provided]')}")
                        percentage = ben.get('percentage', '')
                        doc.add_paragraph(f"  Percentage: {f'{percentage}%' if percentage else '[Not provided]'}")
                else:
                    doc.add_paragraph("[No beneficiaries specified]")
                doc.add_paragraph()
                
                # Asset Breakdown
                doc.add_heading('Asset Breakdown', level=1)
                asset_types = [
                    "Stocks", "Bonds", "Mutual Funds", "ETFs", "UITs", 
                    "Annuities (Fixed)", "Annuities (Variable)", "Options", 
                    "Commodities", "Alternative Investments", "Limited Partnerships", 
                    "Variable Contracts", "Short-Term", "Other"
                ]
                for asset_type in asset_types:
                    field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}"
                    value = self.form_data.get(field_name)
                    doc.add_paragraph(f"{asset_type}: {f'{value}%' if value else '[Not provided]'}")
                doc.add_paragraph()
                
                # Investment Experience
                doc.add_heading('Investment Experience', level=1)
                experience_types = [
                    "Stocks", "Bonds", "Mutual Funds", "UITs", 
                    "Annuities (Fixed)", "Annuities (Variable)", "Options", 
                    "Commodities", "Alternative Investments", "Limited Partnerships", 
                    "Variable Contracts"
                ]
                for exp_type in experience_types:
                    doc.add_paragraph(f"{exp_type}:")
                    year_field = f"asset_experience_{exp_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}_year"
                    level_field = f"asset_experience_{exp_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}_level"
                    year = self.form_data.get(year_field)
                    level = self.form_data.get(level_field)
                    doc.add_paragraph(f"  Year Started: {year or '[Not provided]'}")
                    doc.add_paragraph(f"  Experience Level: {level or '[Not provided]'}")
                doc.add_paragraph()
                
                # Outside Broker Information
                if self.form_data.get('has_outside_broker'):
                    doc.add_heading('Outside Broker Information', level=1)
                    doc.add_paragraph(f"Broker Firm Name: {self.form_data.get('outside_firm_name', '[Not provided]')}")
                    doc.add_paragraph(f"Account Type: {self.form_data.get('outside_broker_account_type', '[Not provided]')}")
                    doc.add_paragraph(f"Account Number: {self.form_data.get('outside_broker_account_number', '[Not provided]')}")
                    doc.add_paragraph(f"Liquid Amount: {format_money(self.form_data.get('outside_liquid_amount'))}")
                    doc.add_paragraph()
                
                # Trusted Contact Information
                doc.add_heading('Trusted Contact Information', level=1)
                doc.add_paragraph(f"Full Name: {self.form_data.get('trusted_full_name', '[Not provided]')}")
                doc.add_paragraph(f"Relationship: {self.form_data.get('trusted_relationship', '[Not provided]')}")
                doc.add_paragraph(f"Phone Number: {self.form_data.get('trusted_phone', '[Not provided]')}")
                doc.add_paragraph(f"Email Address: {self.form_data.get('trusted_email', '[Not provided]')}")
                doc.add_paragraph()
                
                # Regulatory Consent
                doc.add_heading('Regulatory Consent', level=1)
                electronic_consent = "Yes" if self.form_data.get('electronic_regulatory_yes') else "No"
                doc.add_paragraph(f"Electronic Delivery Consent: {electronic_consent}")
                
                # Save the document
                doc.save(file_path)
                QMessageBox.information(self, "Success", "Draft saved successfully in Word format!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save draft: {e}")
                
    def generate_pdf_report(self):
        """Generate a PDF report from the form data"""
        try:
            # Get the output directory
            output_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # Get the file path
            file_path = os.path.join(output_dir, "Magnus_Client_Intake_Form.pdf")
            
            # Generate the PDF
            from pdf_generator_reportlab import generate_pdf_report
            if generate_pdf_report(self.form_data, file_path):
                QMessageBox.information(
                    self,
                    "Success",
                    f"PDF has been generated successfully and saved to:\n{file_path}"
                )
                # Try to open the PDF
                try:
                    import subprocess
                    if os.name == 'nt':  # Windows
                        os.startfile(file_path)
                    elif os.name == 'posix':  # macOS or Linux
                        subprocess.run(['open' if sys.platform == 'darwin' else 'xdg-open', file_path])
                except Exception as e:
                    print(f"Error opening PDF: {str(e)}")
            else:
                QMessageBox.critical(
                    self,
                    "Error",
                    "Failed to generate PDF. Please check the console for details."
                )
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"An error occurred while generating the PDF:\n{str(e)}"
            )
            print(f"Error generating PDF: {str(e)}")
            traceback.print_exc()



def main():
    """Main application entry point"""
    app = QApplication(sys.argv)
    app.setApplicationName("Magnus Client Intake Form")
    app.setApplicationVersion("2.2")
    
    # Set application style
    app.setStyleSheet("""
        QMainWindow {
            background-color: #ffffff;
        }
        QLabel {
            color: #2c3e50;
        }
        QGroupBox {
            font-weight: bold;
            border: 2px solid #bdc3c7;
            border-radius: 5px;
            margin-top: 10px;
            padding-top: 10px;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 5px 0 5px;
        }
    """)
    
    # Create and show main window
    window = MagnusClientIntakeForm()
    window.show()
    
    return app.exec()


if __name__ == "__main__":
    sys.exit(main())


